from typing import override
from docx.document import Document
from docx.text.paragraph import Paragraph
from pathlib import Path
from model.fillpdf.document_template import DocumentTemplate, DocumentWithNoInputs
import re, docx, io, tempfile, os#, docx2pdf



class DOCXTemplate(DocumentTemplate):
    def __init__(self, file_path) -> None:
        super().__init__(file_path)

        self.file: io.BufferedReader = open(file_path, "rb")

        self.binary_data = io.BytesIO(self.file.read())

        self.document: Document = docx.Document(self.binary_data)

        paragraphs_with_input: list[int] = []

        for index in range(len(self.document.paragraphs)):
            text: str = self.document.paragraphs[index].text
            inputs_found: list[str] = re.findall("\\[.*?\\]",text)
            if len(inputs_found) > 0:
                paragraphs_with_input.append(index)
                for input in inputs_found:
                    self.valid_inputs.add(input)
        
        if len(self.valid_inputs) == 0:
            raise DocumentWithNoInputs
        
        self.paragraphs_with_input: tuple[str,...] = tuple(paragraphs_with_input)

        self.temporary_folder:tempfile.TemporaryDirectory = None
    
    @override
    def set_export_type(self, file_type) -> None:
        super().set_export_type(file_type)
        match file_type:
            case "docx":
                self._current_export_function = self.export_docx
            case "pdf":
                self._current_export_function = self.export_pdf
            case _:
                raise NotImplementedError("'%s' file type export is not implemented!")
    
    def replace_inputs(self, headers: tuple[str, ...], data: tuple[str, ...]) -> Document:
        export_document: Document = docx.Document(self.binary_data)

        for index in self.paragraphs_with_input:
            paragraph: Paragraph = export_document.paragraphs[index]
            new_text: str = paragraph.text
            for i in range(len(headers)):
                new_text = new_text.replace(headers[i],data[i])
            paragraph.text = new_text
            export_document.paragraphs[index] = paragraph
        
        return export_document

    def export_docx(self, headers: tuple[str, ...], data: tuple[str, ...], path: Path):
        super().export(headers, data, path)
        export_document: Document = self.replace_inputs(headers, data)
        
        export_document.save(path)
    
    def export_pdf(self, headers: tuple[str, ...], data: tuple[str, ...], path: Path) -> None:
        super().export(headers, data, path)

        export_document: Document = self.replace_inputs(headers, data)
        docx_file_path = os.path.join(self.temporary_folder.name, path.name)
        pdf_file_path = path.parent / (path.stem + ".pdf")

        print(docx_file_path)
        print(pdf_file_path)

        export_document.save(docx_file_path)

        # docx2pdf.convert(docx_file_path, pdf_file_path)


    @override
    def prepare_export(self) -> None:
        if self.export_type == "docx": return

        self.temporary_folder = tempfile.TemporaryDirectory(ignore_cleanup_errors=True)

    @override
    def clean_export(self) -> None:
        if self.export_type == "docx": return
        print("Cleaning...")
        self.temporary_folder.cleanup()
        